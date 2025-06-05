from dotenv import load_dotenv
import os
import streamlit as st
import pyperclip
import pandas as pd
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Load environment variables from .env file
load_dotenv()

os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = './sqy-prod.json'

# Initialize Gemini client
gemini_client = genai.Client(
    http_options=types.HttpOptions(api_version="v1beta1"),
    vertexai=True,
    project='sqy-prod',
    location='us-central1'
)

gemini_tools = [types.Tool(google_search=types.GoogleSearch())]

# --- Load CSV Data Function ---
@st.cache_data
def load_csv_data(csv_file):
    """Load and cache CSV data"""
    try:
        df = pd.read_csv(csv_file)
        # Clean column names (remove extra spaces)
        df.columns = df.columns.str.strip()
        # Ensure we have the expected columns
        expected_columns = ['CItyName', 'micromarket', 'locality']
        if not all(col in df.columns for col in expected_columns):
            st.error(f"CSV file must contain columns: {expected_columns}")
            return None
        return df
    except Exception as e:
        st.error(f"Error loading CSV file: {str(e)}")
        return None

def get_cities_from_csv(df):
    """Get unique cities from CSV"""
    if df is not None:
        return sorted(df['CItyName'].unique().tolist())
    return []

def get_micromarkets_from_csv(df, city):
    """Get micromarkets for a specific city"""
    if df is not None and city:
        micromarkets = df[df['CItyName'] == city]['micromarket'].unique().tolist()
        return sorted(micromarkets)
    return []

def get_localities_from_csv(df, city, micromarket):
    """Get localities for a specific city and micromarket"""
    if df is not None and city and micromarket:
        localities = df[(df['CItyName'] == city) & (df['micromarket'] == micromarket)]['locality'].unique().tolist()
        # Remove NaN values and empty strings
        localities = [loc for loc in localities if pd.notna(loc) and str(loc).strip()]
        return sorted(localities)
    return []

# --- City Description Function ---
def create_city_description(prompt: str, city: str) -> str:
    try:
        full_query = prompt.format(city=city)
        response = gemini_client.models.generate_content(
            model="gemini-2.0-flash-001",
            contents=full_query,
            config=types.GenerateContentConfig(
                tools=gemini_tools,
                max_output_tokens=8192,
                system_instruction="You are a helpful real-estate agent. Provide a response in plain text with Markdown syntax for bold headings (e.g., **Heading**) and no HTML tags, special characters, or FAQs. Do not include suggestions or notes in the response. Use data from Google Search tools to ensure accuracy and relevance.",
                temperature=0.7,
            )
        )
        content = response.text
        return content.strip()
    except Exception as e:
        return f"Error generating city description: {str(e)}"

# --- Enhanced Micro Market Description Function with Localities ---
def create_micromarket_description(prompt: str, city: str, micromarket: str, localities: list = None) -> str:
    try:
        # Format localities for the prompt
        localities_text = ""
        if localities and len(localities) > 0:
            localities_list = [loc.strip() for loc in localities if loc.strip()]
            if localities_list:
                localities_text = f"\nLocalities to focus on: {', '.join(localities_list)}"
        
        full_query = prompt.format(city=city, micromarket=micromarket, localities=localities_text)
        response = gemini_client.models.generate_content(
            model="gemini-2.0-flash-001",
            contents=full_query,
            config=types.GenerateContentConfig(
                tools=gemini_tools,
                max_output_tokens=8192,
                system_instruction="You are a helpful real-estate agent. Provide a response in plain text with Markdown syntax for bold headings (e.g., **Heading**) and no HTML tags, special characters, or FAQs. Do not include suggestions or notes in the response. Use data from Google Search tools to ensure accuracy and relevance.",
                temperature=0.7,
            )
        )
        content = response.text
        return content.strip()
    except Exception as e:
        return f"Error generating micro market description: {str(e)}"

# --- Function to Create DOCX File ---
def create_docx(city_description: str = "", micromarket_description: str = "", city: str = "", micromarket: str = ""):
    doc = Document()
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Combine descriptions (city first, then micro market)
    combined_text = city_description
    if city_description and micromarket_description:
        combined_text += "\n\n" + micromarket_description
    elif micromarket_description:
        combined_text = micromarket_description

    if not combined_text:
        doc.add_paragraph("No descriptions available.")
    else:
        # Split combined text into lines
        lines = combined_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if the line is a heading (starts and ends with **)
            if line.startswith('**') and line.endswith('**'):
                # Add heading with bold formatting
                heading_text = line[2:-2].strip()  # Remove ** markers
                p = doc.add_paragraph()
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(14)  # Larger for headings
            else:
                # Add regular text or bullet points
                if line.startswith('- '):
                    # Bullet point
                    bullet_text = line[2:].strip()  # Remove '- '
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(bullet_text)
                else:
                    # Regular paragraph
                    p = doc.add_paragraph()
                    run = p.add_run(line)

    # Save to BytesIO for download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit App ---
def main():
    st.title("City and Micro Market Description Generator")
    st.write("Generate city and micro market descriptions individually and download them combined in a single DOCX file.")

    # CSV File Upload Section
    st.header("Upload CSV Data")
    uploaded_file = st.file_uploader(
        "Upload CSV file with City, Micromarket, and Locality data",
        type=['csv'],
        help="CSV should have columns: CItyName, micromarket, locality"
    )
    
    csv_data = None
    if uploaded_file is not None:
        csv_data = load_csv_data(uploaded_file)
        if csv_data is not None:
            st.success(f"✅ CSV loaded successfully! Found {len(csv_data)} records")
            
            # Show a preview of the data
            with st.expander("Preview CSV Data"):
                st.dataframe(csv_data.head(10))
                st.write(f"**Total Records:** {len(csv_data)}")
                st.write(f"**Cities:** {len(csv_data['CItyName'].unique())}")
                st.write(f"**Micromarkets:** {len(csv_data['micromarket'].unique())}")
    else:
        st.info("Please upload a CSV file to enable automatic locality selection")

    # City description prompt
    city_prompt = """
    Provide a comprehensive description for city {city}. 
    Use data from Google Search tools to ensure accuracy, including relevant pincode details or other trending information where applicable.
    The description should cover:
    **{city} City Description**
    **Introduction**
    - Brief introduction to the city, including its location, significance, and key characteristics.
    **History**
    - Historical background, including key events or developments that shaped the city.
    **Economy**
    - Economic overview, highlighting major industries, business hubs, and economic challenges.
    **Demography**
    - Demographic details, including population, gender ratio, literacy rate, and notable demographic trends.
    **Infrastructure**
    - Highways: Major highways and roads ensuring connectivity.
    - Metro Routes: Key metro lines and stations for urban connectivity.
    - Rail Routes: Major railway stations and their connectivity.
    - Airport: Details of the nearest airport and its accessibility.
    **Top 20 Builders in {city}**
    - List the top 20 real estate builders operating in the city.
    **Top 10 Schools in {city}**
    - List the top 10 schools in the city.
    **Top 10 Hospitals in {city}**
    - List the top 10 hospitals in the city.
    **Top 10 Malls in {city}**
    - List the top 10 shopping malls in the city.
    Response Format
    - Use plain text with Markdown syntax for bold headings (e.g., **{city} City Description**, **Introduction**) and no HTML tags, special characters, or FAQs.
    - Structure the response with bold section headings using ** and regular text for other content (paragraphs, bullet points).
    - Use bullet points (denoted by '-') for lists.
    - Ensure proper spacing between sections and list items.
    - Use simple, natural, realistic language without embellishment.
    - Do not add notes, FAQs, or extra commentary.
    """

    # Enhanced micro market description prompt with localities
    micromarket_prompt = """
    Provide a comprehensive description for micro market {micromarket} in city {city}. 
    Use data from Google Search tools to ensure accuracy, including relevant pincode details or other trending information where applicable.{localities}
    The description should cover:
    **{micromarket} Micro Market Description**
    - Provide a description covering:
      - Unique selling points (connectivity, amenities, property rates, lifestyle).
      - Location within the city (describe relative to major roads, metro stations, or landmarks).
      - History of {micromarket} with respect to real estate and infrastructure development.
    **Key Areas in {micromarket}**
    - List the key areas in {micromarket}, each with their unique selling points (connectivity, amenities, property rates, lifestyle).
    **Locality-wise Analysis**
    - If specific localities are mentioned, provide detailed analysis for each locality including:
      - Connectivity and transportation options
      - Property price trends and types available
      - Nearby amenities (schools, hospitals, malls, parks)
      - Infrastructure development and future projects
      - Residential vs commercial mix
      - Target demographics and lifestyle
    **Investment Potential**
    - Analyze the investment potential of {micromarket} and its localities
    - Price appreciation trends and future growth prospects
    - Rental yields and demand patterns
    **Infrastructure Development**
    - Current and upcoming infrastructure projects
    - Government initiatives and policy impacts
    - Transportation connectivity improvements
    Response Format
    - Use plain text with Markdown syntax for bold headings (e.g., **{micromarket} Micro Market Description**) and no HTML tags, special characters, or FAQs.
    - Structure the response with bold section headings using ** and regular text for other content (paragraphs, bullet points).
    - Use bullet points (denoted by '-') for lists.
    - Ensure proper spacing between sections and list items.
    - Use simple, natural, realistic language without embellishment.
    - Do not add notes, FAQs, or extra commentary.
    """

    # City Description Form
    st.header("Generate City Description")
    with st.form(key='city_form'):
        if csv_data is not None:
            # Dropdown for city selection from CSV
            available_cities = get_cities_from_csv(csv_data)
            city = st.selectbox("Select City", [""] + available_cities, key="city_select")
        else:
            # Text input if no CSV is uploaded
            city = st.text_input("City (e.g., Gurgaon)", key="city_input")
        
        city_prompt_area = st.text_area("Edit City Description Prompt", city_prompt, height=300, key="city_prompt")
        city_submit = st.form_submit_button(label='Generate City Description')

    if city_submit and city:
        with st.spinner("Generating city description..."):
            city_description = create_city_description(city_prompt_area, city)
            st.session_state['city_description'] = city_description
            st.session_state['city_name'] = city

    # Display city description if available
    if 'city_description' in st.session_state:
        st.markdown("### City Description")
        st.markdown(st.session_state['city_description'])
        if st.button("Copy City Description"):
            try:
                pyperclip.copy(st.session_state['city_description'])
                st.success("City description copied to clipboard!")
            except Exception as e:
                st.error(f"Failed to copy city description: {str(e)}")

    # Enhanced Micro Market Description Form with CSV-based Localities
    st.header("Generate Micro Market Description")
    
    if csv_data is not None:
        # Outside the form for dynamic updates
        available_cities = get_cities_from_csv(csv_data)
        micro_city = st.selectbox("Select City", [""] + available_cities, key="micro_city_select")
        
        # Dropdown for micromarket selection based on selected city
        available_micromarkets = []
        if micro_city:
            available_micromarkets = get_micromarkets_from_csv(csv_data, micro_city)
        
        micromarket = st.selectbox(
            "Select Micro Market", 
            [""] + available_micromarkets, 
            key="micromarket_select",
            disabled=not micro_city
        )
        
        # Auto-populate localities based on city and micromarket selection
        localities = []
        if micro_city and micromarket:
            localities = get_localities_from_csv(csv_data, micro_city, micromarket)
            
        # Display auto-filled localities
        if localities:
            st.success(f"✅ Found {len(localities)} localities for {micromarket}, {micro_city}")
            with st.expander("View Auto-filled Localities"):
                for i, locality in enumerate(localities, 1):
                    st.write(f"{i}. {locality}")
        elif micro_city and micromarket:
            st.warning("No localities found for the selected combination")
    
    # Form only contains the prompt and submit button
    with st.form(key='micromarket_form'):
        if csv_data is None:
            # Manual input if no CSV is uploaded
            micro_city = st.text_input("City (e.g., Gurgaon)", key="micro_city_input")
            micromarket = st.text_input("Micro Market (e.g., Golf Course Road)", key="micromarket_input")
            
            # Manual locality input
            st.write("Enter localities separated by commas (e.g., Sector 54, DLF Phase 3, Cyber Hub)")
            locality_text = st.text_area(
                "Localities (one per line or comma-separated):",
                placeholder="Enter localities like:\nSector 54\nDLF Phase 3\nCyber Hub\n\nOr comma-separated: Sector 54, DLF Phase 3, Cyber Hub",
                key="locality_text"
            )
            
            localities = []
            if locality_text:
                # Handle both comma-separated and line-separated input
                if ',' in locality_text:
                    localities = [loc.strip() for loc in locality_text.split(',') if loc.strip()]
                else:
                    localities = [loc.strip() for loc in locality_text.split('\n') if loc.strip()]
        
        micromarket_prompt_area = st.text_area("Edit Micro Market Description Prompt", micromarket_prompt, height=300, key="micromarket_prompt")
        micromarket_submit = st.form_submit_button(label='Generate Micro Market Description')

    if micromarket_submit and micro_city and micromarket:
        with st.spinner("Generating micro market description..."):
            micromarket_description = create_micromarket_description(micromarket_prompt_area, micro_city, micromarket, localities)
            st.session_state['micromarket_description'] = micromarket_description
            st.session_state['micromarket_name'] = micromarket
            st.session_state['micro_city_name'] = micro_city
            st.session_state['localities'] = localities

    # Display micro market description if available
    if 'micromarket_description' in st.session_state:
        st.markdown("### Micro Market Description")
        
        # Show selected localities info
        if 'localities' in st.session_state and st.session_state['localities']:
            st.info(f"Generated for {len(st.session_state['localities'])} localities: {', '.join(st.session_state['localities'][:5])}{'...' if len(st.session_state['localities']) > 5 else ''}")
        
        st.markdown(st.session_state['micromarket_description'])
        if st.button("Copy Micro Market Description"):
            try:
                pyperclip.copy(st.session_state['micromarket_description'])
                st.success("Micro market description copied to clipboard!")
            except Exception as e:
                st.error(f"Failed to copy micro market description: {str(e)}")

    # Download Combined DOCX
    if ('city_description' in st.session_state or 'micromarket_description' in st.session_state):
        st.header("Download Combined Description")
        city_desc = st.session_state.get('city_description', "")
        micro_desc = st.session_state.get('micromarket_description', "")
        city_name = st.session_state.get('city_name', "city")
        micro_name = st.session_state.get('micromarket_name', "micromarket")
        if st.button("Download Combined DOCX"):
            docx_buffer = create_docx(city_desc, micro_desc, city_name, micro_name)
            st.download_button(
                label="Download Combined Description as DOCX",
                data=docx_buffer,
                file_name=f"{micro_name}_{city_name}_description.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()